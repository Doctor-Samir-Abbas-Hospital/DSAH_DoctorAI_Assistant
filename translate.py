import os
from dotenv import load_dotenv
import streamlit as st
from langchain_core.messages import AIMessage
from PyPDF2 import PdfReader, PdfWriter
from io import BytesIO
import arabic_reshaper
from bidi.algorithm import get_display
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.lib.fonts import addMapping
from reportlab.lib.utils import simpleSplit
from utils.functions import get_vector_store, get_response_
from spire.pdf.common import *
from spire.pdf import PdfDocument, FileFormat
from docx import Document
from datetime import datetime

# Load environment variables
load_dotenv()

# Register Arabic font (Arial)
font_path = os.path.join('assets', 'Arial.ttf')
pdfmetrics.registerFont(TTFont('Arial', font_path))
addMapping('Arial', 0, 0, 'Arial')

# Path to the GIF image
gif_path = os.path.join('assets', 'tap.gif')

def reshape_arabic_text(text):
    """Reshapes and applies bidi formatting for Arabic text."""
    reshaped_text = arabic_reshaper.reshape(text)
    bidi_text = get_display(reshaped_text)
    return bidi_text

def clean_text(text):
    """Removes unwanted characters and cleans the text."""
    return text.replace('*', '').replace('#', '')

def create_pdf(translated_text, doctor_name, department, selected_date):
    """Creates a PDF with translated text and additional info."""
    try:
        template_path = os.path.join('assets', 'hospital_template.pdf')
        if not os.path.exists(template_path):
            raise FileNotFoundError("Hospital template PDF not found")

        text_buffer = BytesIO()
        c = canvas.Canvas(text_buffer, pagesize=A4)
        c.setTitle("Translated Medical Report")
        c.setAuthor("Doctor AI Assistant")

        width, height = A4
        margin = inch
        text_width = width - 2 * margin

        reshaped_text = reshape_arabic_text(clean_text(translated_text))
        lines = simpleSplit(reshaped_text, 'Arial', 12, text_width)

        y = height - 2 * margin
        c.setFont("Arial", 12)

        for line in lines:
            if y < 1.8 * margin:
                c.showPage()
                c.setFont("Arial", 12)
                y = height - 2 * margin

            c.drawRightString(width - margin, y, line)
            y -= 14

        # Add doctor details below the translation
        if y < margin:
            c.showPage()
            y = height - 2 * margin

        y -= 20
        c.setFont("Arial", 12)
        c.drawRightString(width - margin, y, f"Doctor: {doctor_name}")
        y -= 14
        c.drawRightString(width - margin, y, f"Department: {department}")
        y -= 14
        c.drawRightString(width - margin, y, f"Date: {selected_date}")

        c.save()
        text_buffer.seek(0)

        output_buffer = BytesIO()
        template_pdf = PdfReader(template_path)
        output_pdf = PdfWriter()
        text_pdf = PdfReader(text_buffer)

        for i in range(max(len(template_pdf.pages), len(text_pdf.pages))):
            template_page = template_pdf.pages[0 if len(template_pdf.pages) == 1 else i % len(template_pdf.pages)]
            new_page = PdfWriter().add_blank_page(width=template_page.mediabox.width, 
                                                  height=template_page.mediabox.height)
            new_page.merge_page(template_page)

            if i < len(text_pdf.pages):
                text_page = text_pdf.pages[i]
                new_page.merge_page(text_page)

            output_pdf.add_page(new_page)

        output_buffer = BytesIO()
        output_pdf.write(output_buffer)
        output_buffer.seek(0)
        
        return output_buffer

    except Exception as e:
        st.error(f"Error creating PDF: {str(e)}")
        return None

def convert_pdf_to_word(pdf_path, output_path="output.docx"):
    """Convert PDF to Word using Spire.PDF."""
    try:
        doc = PdfDocument()
        doc.LoadFromFile(pdf_path)
        doc.SaveToFile(output_path, FileFormat.DOCX)
        doc.Close()
        
        with open(output_path, "rb") as file:
            word_buffer = BytesIO(file.read())
        word_buffer.seek(0)
        
        return word_buffer

    except Exception as e:
        st.error(f"Error converting PDF to Word: {str(e)}")
        return None

def read_docx(file):
    """Read a DOCX file and return its text."""
    doc = Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def translate():
    if "translate_state" not in st.session_state:
        st.session_state.translate_state = {}

    if "chat_history1" not in st.session_state:
        st.session_state.chat_history1 = []

    if "vector_store" not in st.session_state:
        st.session_state.vector_store = get_vector_store()

    if "pdf_text" not in st.session_state:
        st.session_state.pdf_text = ""

    if "translated_text" not in st.session_state:
        st.session_state.translated_text = ""

    with open('style.css') as f:
        st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)

    with st.sidebar:
        title = "Doctor AI Assistant"
        name = "Medical Diagnosis and More..."
        profession = "Doctor Samir Abbas Hospital"
        imgUrl = "https://media3.giphy.com/media/6P47BlxlgrJxQ9GR58/giphy.gif"

        st.markdown(
            f"""
                <img class="profileImage" src="{imgUrl}" alt="Your Photo">
                <div class="textContainer">
                    <div class="title"><p>{title}</p></div>
                    <p>{name}</p>
                    <p>{profession}</p>
                    <p>Powered by DSAH Information Technology</p>
                </div>
            """,
            unsafe_allow_html=True,
        )

        doctor_name = st.text_input("Doctor Name")
        department = st.text_input("Department")
        selected_date = st.date_input("Select Date", datetime.now())

    if doctor_name and department and selected_date:
        uploaded_file = st.sidebar.file_uploader("Upload a medical report", type=["pdf", "docx", "txt"])
        st.session_state.uploaded_file = uploaded_file
    
        if uploaded_file:
            # Check if the translate button has been pressed
            if "translate_pressed" not in st.session_state:
                st.session_state.translate_pressed = False

            col1, col2 = st.columns([1, 2])

            with col1:
                # Show the GIF only if the translate button has not been pressed
                if not st.session_state.translate_pressed:
                    st.image(gif_path, use_column_width=True)

            with col2:
                translate_button = st.button("👉 Press To Translate The Medical Report")
                if translate_button:
                    # Set flag to true to hide the GIF
                    st.session_state.translate_pressed = True
                    st.markdown("""<div class="typewriter"><div class="slide"><i></i></div><div class="paper"></div><div class="keyboard"></div></div>""", unsafe_allow_html=True)

                    with st.spinner("Please wait, it's translating the text..."):
                        if uploaded_file.type == "application/pdf":
                            reader = PdfReader(uploaded_file)
                            st.session_state.pdf_text = ""
                            for page in reader.pages:
                                st.session_state.pdf_text += page.extract_text()
                        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                            st.session_state.pdf_text = read_docx(uploaded_file)
                        elif uploaded_file.type == "text/plain":
                            st.session_state.pdf_text = uploaded_file.read().decode()

                        translation_prompt = "Please translate the attached document comprehensively into medical Arabic in a well-structured format."
                        response = get_response_(translation_prompt + " " + st.session_state.pdf_text)
                        st.session_state.chat_history1.append(AIMessage(content=response))
                        st.session_state.translated_text = clean_text(response)

                    st.markdown("<style>.typewriter { display: none; }</style>", unsafe_allow_html=True)
    else:
        st.sidebar.info("Please fill your name and department press ENTER to upload the documents")

    if st.session_state.translated_text:
        styled_text = st.session_state.translated_text.replace('\n', '</p><p>')
        st.components.v1.html(
            f"""
            <div class='translatedText' data-testid="stAppViewContainer">
                <textarea id="editableText" style="width:100%; height:600px;" readonly>{st.session_state.translated_text}</textarea>
                <button onclick="copyToClipboard()" style="margin-top: 10px; padding: 10px 15px;">Copy Translation 📕</button>
            </div>
            <script>
                function copyToClipboard() {{
                    var copyText = document.getElementById('editableText');
                    copyText.select();
                    document.execCommand("copy");
                    alert("Copied to clipboard!");
                }}
            </script>
            <style>
               [data-testid="stAppViewContainer"] {{
                    background-color: white;
                    overflow-y: auto;
                    max-height: 95vh;
                    width: 100%;
                    direction: rtl;
                    text-align: justify;
                    font-family: Arial, sans-serif;
                    box-sizing: border-box;
                    padding: 20px;
                    border: 1px solid #ccc;
                }}
                textarea {{
                    font-size: 18px;
                    line-height: 1;
                    margin: 0 0 10px 0;
                    border: 1px solid #ddd;
                    border-radius: 4px;
                    padding: 10px;
                }}
                button {{
                    cursor: pointer;
                    background-color: #4CAF50;
                    color: white;
                    border: none;
                    border-radius: 5px;
                }}
            </style>
            """,
            height=800,
            width=800,
        )
        pdf_buffer = create_pdf(
            st.session_state.translated_text,
            doctor_name,
            department,
            selected_date.strftime("%Y-%m-%d")
        )
        
        if pdf_buffer:
            pdf_path = "translated_report.pdf"
            with open(pdf_path, "wb") as f:
                f.write(pdf_buffer.getbuffer())
                
            word_buffer = convert_pdf_to_word(pdf_path)
            
            st.sidebar.download_button(
                label="Download Translated Report as PDF",
                data=pdf_buffer,
                file_name="translated_report.pdf",
                mime="application/pdf"
            )
            
            if word_buffer:
                st.sidebar.download_button(
                    label="Download Translated Report as Word",
                    data=word_buffer,
                    file_name="translated_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.error("Unable to generate Word document.")

if __name__ == "__main__":
    translate()






